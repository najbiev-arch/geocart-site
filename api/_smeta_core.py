from __future__ import annotations

import io
import json
import mimetypes
import re
import sys
from datetime import date, datetime
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import urlparse


ROOT = Path(__file__).resolve().parent
STATIC_DIR = ROOT / "static"


URBAN_WORK_TYPES: list[dict[str, Any]] = [
    {
        "id": "city_general_plan",
        "title": "Генеральный план города или поселка",
        "table": "1706-0102-01",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-01",
        "indicatorLabel": "Численность населения",
        "indicatorUnit": "тыс. чел.",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "5", "label": "от 0,5 до 1 тыс. чел.", "min": 0.5, "max": 1, "a": 7523, "b": 1220},
            {"pos": "6", "label": "свыше 1 до 2 тыс. чел.", "min": 1, "max": 2, "minExclusive": True, "a": 7763, "b": 976},
            {"pos": "7", "label": "свыше 2 до 5 тыс. чел.", "min": 2, "max": 5, "minExclusive": True, "a": 7890, "b": 908},
            {"pos": "8", "label": "свыше 5 до 10 тыс. чел.", "min": 5, "max": 10, "minExclusive": True, "a": 8800, "b": 727},
            {"pos": "9", "label": "свыше 10 до 25 тыс. чел.", "min": 10, "max": 25, "minExclusive": True, "a": 9192, "b": 688},
            {"pos": "10", "label": "свыше 25 до 49,9 тыс. чел.", "min": 25, "max": 49.9, "minExclusive": True, "a": 18156, "b": 329},
            {"pos": "11", "label": "от 50 до 100 тыс. чел.", "min": 50, "max": 100, "a": 28941, "b": 290},
            {"pos": "12", "label": "свыше 100 до 250 тыс. чел.", "min": 100, "max": 250, "minExclusive": True, "a": 33278, "b": 247},
            {"pos": "13", "label": "свыше 250 до 499,9 тыс. чел.", "min": 250, "max": 499.9, "minExclusive": True, "a": 55936, "b": 157},
            {"pos": "14", "label": "от 500 до 750 тыс. чел.", "min": 500, "max": 750, "a": 58401, "b": 115},
            {"pos": "15", "label": "свыше 750 до 1000 тыс. чел.", "min": 750, "max": 1000, "minExclusive": True, "a": 59565, "b": 114},
            {"pos": "16", "label": "свыше 1000 до 1500 тыс. чел.", "min": 1000, "max": 1500, "minExclusive": True, "a": 95713, "b": 78},
            {"pos": "17", "label": "свыше 1500 до 2500 тыс. чел.", "min": 1500, "max": 2500, "minExclusive": True, "a": 120533, "b": 61},
        ],
    },
    {
        "id": "city_general_plan_teo",
        "title": "ТЭО генерального плана города",
        "table": "1706-0102-01",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-01",
        "indicatorLabel": "Численность населения",
        "indicatorUnit": "тыс. чел.",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 500 до 750 тыс. чел.", "min": 500, "max": 750, "a": 26097, "b": 52},
            {"pos": "2", "label": "свыше 750 до 1000 тыс. чел.", "min": 750, "max": 1000, "minExclusive": True, "a": 32014, "b": 44},
            {"pos": "3", "label": "свыше 1000 до 1500 тыс. чел.", "min": 1000, "max": 1500, "minExclusive": True, "a": 42205, "b": 34},
            {"pos": "4", "label": "свыше 1500 до 2500 тыс. чел.", "min": 1500, "max": 2500, "minExclusive": True, "a": 53116, "b": 27},
        ],
    },
    {
        "id": "city_position_scheme_with_materials",
        "title": "Схема положения города в системе расселения, есть материалы районной планировки",
        "table": "1706-0102-02",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-02",
        "indicatorLabel": "Территория пригородной и зеленой зон",
        "indicatorUnit": "10 тыс. га",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 10 до 50 тыс. га", "min": 1, "max": 5, "a": 2895, "b": 1005},
            {"pos": "2", "label": "свыше 50 до 200 тыс. га", "min": 5, "max": 20, "minExclusive": True, "a": 5626, "b": 365},
            {"pos": "3", "label": "свыше 200 до 800 тыс. га", "min": 20, "max": 80, "minExclusive": True, "a": 9091, "b": 216},
        ],
    },
    {
        "id": "city_position_scheme_without_materials",
        "title": "Схема положения города в системе расселения, без материалов районной планировки",
        "table": "1706-0102-02",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-02",
        "indicatorLabel": "Территория пригородной и зеленой зон",
        "indicatorUnit": "10 тыс. га",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "4", "label": "от 10 до 50 тыс. га", "min": 1, "max": 5, "a": 4666, "b": 1601},
            {"pos": "5", "label": "свыше 50 до 200 тыс. га", "min": 5, "max": 20, "minExclusive": True, "a": 9723, "b": 588},
            {"pos": "6", "label": "свыше 200 до 800 тыс. га", "min": 20, "max": 80, "minExclusive": True, "a": 14566, "b": 345},
        ],
    },
    {
        "id": "industrial_zone_plan",
        "title": "Проект планировки промышленной зоны города",
        "table": "1706-0102-03",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-03",
        "indicatorLabel": "Территория",
        "indicatorUnit": "га",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 30 до 50 га", "min": 30, "max": 50, "a": 5133, "b": 167},
            {"pos": "2", "label": "свыше 50 до 100 га", "min": 50, "max": 100, "minExclusive": True, "a": 8307, "b": 104},
            {"pos": "3", "label": "свыше 100 до 250 га", "min": 100, "max": 250, "minExclusive": True, "a": 11923, "b": 67},
            {"pos": "4", "label": "свыше 250 до 500 га", "min": 250, "max": 500, "minExclusive": True, "a": 15552, "b": 53},
            {"pos": "5", "label": "свыше 500 до 1000 га", "min": 500, "max": 1000, "minExclusive": True, "a": 24213, "b": 37},
        ],
    },
    {
        "id": "city_pdp",
        "title": "Проект детальной планировки района",
        "table": "1706-0102-04",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-04",
        "indicatorLabel": "Территория",
        "indicatorUnit": "га",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 30 до 50 га", "min": 30, "max": 50, "a": 3932, "b": 292},
            {"pos": "2", "label": "свыше 50 до 100 га", "min": 50, "max": 100, "minExclusive": True, "a": 10810, "b": 154},
            {"pos": "3", "label": "свыше 100 до 250 га", "min": 100, "max": 250, "minExclusive": True, "a": 17360, "b": 89},
            {"pos": "4", "label": "свыше 250 до 500 га", "min": 250, "max": 500, "minExclusive": True, "a": 20306, "b": 78},
            {"pos": "5", "label": "свыше 500 до 1000 га", "min": 500, "max": 1000, "minExclusive": True, "a": 33569, "b": 52},
        ],
    },
    {
        "id": "city_district_scheme",
        "title": "Схема планировки городского планировочного района",
        "table": "1706-0102-05",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-05",
        "indicatorLabel": "Численность населения",
        "indicatorUnit": "тыс. чел.",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 100 до 250 тыс. чел.", "min": 100, "max": 250, "a": 7005, "b": 52},
            {"pos": "2", "label": "свыше 250 до 500 тыс. чел.", "min": 250, "max": 500, "minExclusive": True, "a": 9356, "b": 42},
        ],
    },
    {
        "id": "housing_civil_section",
        "title": "Раздел \"Жилищно-гражданское строительство\" в составе проекта промышленного предприятия",
        "table": "1706-0102-06",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0102-06",
        "indicatorLabel": "Численность населения",
        "indicatorUnit": "тыс. чел.",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 1 до 2,5 тыс. чел.", "min": 1, "max": 2.5, "a": 6853, "b": 732},
            {"pos": "2", "label": "свыше 2,5 до 5 тыс. чел.", "min": 2.5, "max": 5, "minExclusive": True, "a": 7169, "b": 609},
            {"pos": "3", "label": "свыше 5 до 10 тыс. чел.", "min": 5, "max": 10, "minExclusive": True, "a": 7270, "b": 590},
            {"pos": "4", "label": "свыше 10 до 25 тыс. чел.", "min": 10, "max": 25, "minExclusive": True, "a": 7548, "b": 561},
            {"pos": "5", "label": "свыше 25 до 50 тыс. чел.", "min": 25, "max": 50, "minExclusive": True, "a": 13858, "b": 309},
            {"pos": "6", "label": "свыше 50 до 100 тыс. чел.", "min": 50, "max": 100, "minExclusive": True, "a": 11569, "b": 354},
            {"pos": "7", "label": "свыше 100 до 200 тыс. чел.", "min": 100, "max": 200, "minExclusive": True, "a": 26779, "b": 201},
        ],
    },
    {
        "id": "rural_general_plan",
        "title": "Генеральный план сельского населенного пункта",
        "table": "1706-0107-01",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0107-01",
        "indicatorLabel": "Численность населения",
        "indicatorUnit": "тыс. чел.",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 0,2 до 0,5 тыс. чел.", "min": 0.2, "max": 0.5, "a": 1770, "b": 5743},
            {"pos": "2", "label": "свыше 0,5 до 1 тыс. чел.", "min": 0.5, "max": 1, "minExclusive": True, "a": 2263, "b": 4767},
            {"pos": "3", "label": "свыше 1 до 2 тыс. чел.", "min": 1, "max": 2, "minExclusive": True, "a": 3983, "b": 3066},
            {"pos": "4", "label": "свыше 2 до 5 тыс. чел.", "min": 2, "max": 5, "minExclusive": True, "a": 8699, "b": 703},
            {"pos": "5", "label": "свыше 5 до 10 тыс. чел.", "min": 5, "max": 10, "minExclusive": True, "a": 9407, "b": 568},
        ],
    },
    {
        "id": "agri_enterprise_scheme",
        "title": "Схема планировки территории сельскохозяйственного предприятия",
        "table": "1706-0107-02",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0107-02",
        "indicatorLabel": "Территория",
        "indicatorUnit": "тыс. га",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 1 до 4 тыс. га", "min": 1, "max": 4, "a": 2680, "b": 417},
            {"pos": "2", "label": "свыше 4 до 10 тыс. га", "min": 4, "max": 10, "minExclusive": True, "a": 2921, "b": 365},
            {"pos": "3", "label": "свыше 10 до 20 тыс. га", "min": 10, "max": 20, "minExclusive": True, "a": 4640, "b": 196},
            {"pos": "4", "label": "свыше 20 до 50 тыс. га", "min": 20, "max": 50, "minExclusive": True, "a": 5614, "b": 147},
            {"pos": "5", "label": "свыше 50 до 100 тыс. га", "min": 50, "max": 100, "minExclusive": True, "a": 9268, "b": 73},
            {"pos": "6", "label": "свыше 100 до 200 тыс. га", "min": 100, "max": 200, "minExclusive": True, "a": 11468, "b": 49},
        ],
    },
    {
        "id": "rural_pdp",
        "title": "Проект детальной планировки сельского населенного пункта",
        "table": "1706-0107-03",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0107-03",
        "indicatorLabel": "Территория",
        "indicatorUnit": "га",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "от 15 до 30 га", "min": 15, "max": 30, "a": 3148, "b": 309},
            {"pos": "2", "label": "свыше 30 до 50 га", "min": 30, "max": 50, "minExclusive": True, "a": 4514, "b": 263},
            {"pos": "3", "label": "свыше 50 до 100 га", "min": 50, "max": 100, "minExclusive": True, "a": 9736, "b": 158},
            {"pos": "4", "label": "свыше 100 до 150 га", "min": 100, "max": 150, "minExclusive": True, "a": 12012, "b": 135},
        ],
    },
    {
        "id": "collective_gardens",
        "title": "Проект организации и застройки территории коллективных садов",
        "table": "1706-0107-05",
        "source": "СЦП РК 8.03-01-2025, раздел 6, таблица 1706-0107-05",
        "indicatorLabel": "Количество участков",
        "indicatorUnit": "уч.",
        "priceUnit": "тыс. тенге",
        "rows": [
            {"pos": "1", "label": "до 20 шт.", "min": 0, "max": 20, "a": 2453, "b": 0, "fixed": True},
            {"pos": "2", "label": "свыше 20 до 50 шт.", "min": 20, "max": 50, "minExclusive": True, "a": 1454, "b": 48},
            {"pos": "3", "label": "свыше 50 до 100 шт.", "min": 50, "max": 100, "minExclusive": True, "a": 2187, "b": 39},
            {"pos": "4", "label": "свыше 100 до 200 шт.", "min": 100, "max": 200, "minExclusive": True, "a": 4400, "b": 18},
            {"pos": "5", "label": "свыше 200 до 300 шт.", "min": 200, "max": 300, "minExclusive": True, "a": 4880, "b": 14},
            {"pos": "6", "label": "свыше 300 до 400 шт.", "min": 300, "max": 400, "minExclusive": True, "a": 7080, "b": 7.59},
        ],
    },
]


TERRITORY_LABELS = {
    "undeveloped": "незастроенная территория",
    "built": "застроенная территория",
    "industrial": "территория действующих промышленных предприятий",
}

WORK_LABELS = {
    "field": "полевые работы",
    "office": "камеральные работы",
}


_TOPO_CACHE: list[dict[str, Any]] | None = None
_TOPO_SOURCE: str | None = None


def find_pdf(name_part: str) -> Path | None:
    lowered = name_part.lower()
    for path in ROOT.rglob("*.pdf"):
        if lowered in path.name.lower() and "рус" in str(path).lower():
            return path
    return None


def parse_topo_rates() -> list[dict[str, Any]]:
    global _TOPO_CACHE, _TOPO_SOURCE
    if _TOPO_CACHE is not None:
        return _TOPO_CACHE

    pdf_path = find_pdf("Раздел 1 Инженерно-геодезические")
    if not pdf_path:
        _TOPO_SOURCE = None
        _TOPO_CACHE = []
        return _TOPO_CACHE

    try:
        from pypdf import PdfReader
    except Exception:
        _TOPO_SOURCE = str(pdf_path)
        _TOPO_CACHE = []
        return _TOPO_CACHE

    reader = PdfReader(str(pdf_path))
    current_table: str | None = None
    lines: list[dict[str, Any]] = []

    for page_no in range(21, 35):
        text = reader.pages[page_no - 1].extract_text(extraction_mode="layout") or ""
        for raw in text.splitlines():
            clean = " ".join(raw.strip().split())
            if not clean:
                continue
            if "Таблица 1601-0102-01" in clean:
                current_table = "1601-0102-01"
            if "Таблица 1601-0102-02" in clean:
                current_table = "1601-0102-02"
            if not current_table:
                continue
            if any(
                marker in clean
                for marker in [
                    "СЦИ РК",
                    "Продолжение таблицы",
                    "Окончание таблицы",
                    "№ позиции",
                    "Наименование работ",
                    "Ед. изм",
                    "Цена, тенге",
                ]
            ):
                continue
            lines.append({"page": page_no, "table": current_table, "raw": raw.rstrip(), "clean": clean})

    hit_re = re.compile(r"^\s*(\d{2})\s+.*\bга\b\s+([0-9]+(?:\s+[0-9]{3})*)\s*$")
    hits: list[tuple[int, str, int]] = []
    for idx, line in enumerate(lines):
        match = hit_re.match(line["raw"])
        if match:
            hits.append((idx, match.group(1), int(match.group(2).replace(" ", ""))))

    rates: list[dict[str, Any]] = []
    for hit_no, (idx, pos, price) in enumerate(hits):
        prev_hit = hits[hit_no - 1][0] if hit_no else -1
        next_hit = hits[hit_no + 1][0] if hit_no + 1 < len(hits) else len(lines)

        start = idx
        for cursor in range(idx, prev_hit, -1):
            clean = lines[cursor]["clean"]
            if "Топографическая съемка" in clean or "Создание инженерно" in clean:
                start = cursor
                break

        end = next_hit
        for cursor in range(idx + 1, next_hit):
            clean = lines[cursor]["clean"]
            if "Топографическая съемка" in clean or "Создание инженерно" in clean or "Примечания" in clean:
                end = cursor
                break

        raw_text = " ".join(lines[cursor]["clean"] for cursor in range(start, end))
        text = re.sub(r"\b" + re.escape(pos) + r"\b", " ", raw_text)
        text = re.sub(r"\bга\b\s+[0-9]+(?:\s+[0-9]{3})*", " ", text)
        text = " ".join(text.split())
        lower = text.lower()

        scale = re.search(r"масштаб\s+съемки\s+1:(\d+)", lower)
        relief = re.search(r"высота\s+сечения\s+рельефа\s+([0-9]+(?:[,.][0-9]+)?)\s*м", lower)
        category = re.search(r"\b(I{1,3})\s+категори", text)

        if "действующих" in lower and "промышлен" in lower:
            territory = "industrial"
        elif "незастроенной" in lower:
            territory = "undeveloped"
        elif "застроенной" in lower:
            territory = "built"
        else:
            territory = "unknown"

        if "полев" in lower:
            work = "field"
        elif "камеральн" in lower:
            work = "office"
        else:
            work = "unknown"

        if not scale or not relief or not category or territory == "unknown" or work == "unknown":
            continue

        rate = {
            "id": f"{lines[idx]['table']}-{pos}",
            "table": lines[idx]["table"],
            "position": pos,
            "page": lines[idx]["page"],
            "scale": int(scale.group(1)),
            "relief": float(relief.group(1).replace(",", ".")),
            "category": category.group(1),
            "territory": territory,
            "territoryLabel": TERRITORY_LABELS[territory],
            "work": work,
            "workLabel": WORK_LABELS[work],
            "unit": "га",
            "price": price,
            "source": f"СЦИ РК 8.03-04-2025, таблица {lines[idx]['table']}, поз. {pos}",
            "description": text,
        }
        rates.append(rate)

    _TOPO_SOURCE = str(pdf_path)
    _TOPO_CACHE = rates
    return rates


def build_catalog() -> dict[str, Any]:
    return {
        "urban": URBAN_WORK_TYPES,
        "topoRates": parse_topo_rates(),
        "topoSource": _TOPO_SOURCE,
        "territoryLabels": TERRITORY_LABELS,
        "workLabels": WORK_LABELS,
        "generatedAt": datetime.now().isoformat(timespec="seconds"),
    }


def read_payload(handler: SimpleHTTPRequestHandler) -> dict[str, Any]:
    length = int(handler.headers.get("Content-Length", "0"))
    raw = handler.rfile.read(length)
    if not raw:
        return {}
    return json.loads(raw.decode("utf-8"))


def clean_filename(value: str) -> str:
    safe = re.sub(r"[^0-9A-Za-zА-Яа-я._ -]+", "_", value).strip(" ._")
    return safe or "smeta"


def estimate_total(items: list[dict[str, Any]]) -> float:
    return round(sum(float(item.get("total", 0) or 0) for item in items), 2)


def normalize_estimate_number(value: str | None) -> str:
    number = (value or "").strip()
    number = re.sub(r"^\s*смета\s*[-№#]*\s*", "", number, flags=re.IGNORECASE)
    return number or "1"


def estimate_title(meta: dict[str, Any]) -> str:
    return f"СМЕТА № {normalize_estimate_number(meta.get('estimateNo'))}"


def estimate_subtitle(meta: dict[str, Any]) -> str:
    return (meta.get("estimateSubtitle") or "на проектные работы").strip()


def estimate_header_rows(meta: dict[str, Any]) -> list[tuple[str, str]]:
    return [
        ("Наименование проектных работ", meta.get("objectName", "")),
        ("Стадия проектирования, этапа, вида, проектных или изыскательских работ", meta.get("stageName", "")),
        ("Наименование проектной (изыскательской) организации", meta.get("designer", "")),
        ("Наименование организации заказчика", meta.get("customer", "")),
    ]


def create_docx(payload: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    meta = payload.get("meta", {})
    items = payload.get("items", [])
    total = estimate_total(items)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.2)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    def set_cell(cell: Any, text: str, *, bold: bool = False, center: bool = False, size: int = 10) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text or "")
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def clear_cell_borders(cell: Any) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            element.set(qn("w:val"), "nil")

    def clear_table_borders(table: Any) -> None:
        seen: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id not in seen:
                    seen.add(cell_id)
                    clear_cell_borders(cell)

    header_table = document.add_table(rows=6, cols=5)
    header_table.style = "Table Grid"
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    title_cell = header_table.rows[0].cells[0].merge(header_table.rows[0].cells[-1])
    set_cell(title_cell, estimate_title(meta), bold=True, center=True, size=12)
    subtitle_cell = header_table.rows[1].cells[0].merge(header_table.rows[1].cells[-1])
    set_cell(subtitle_cell, estimate_subtitle(meta), bold=True, center=True, size=11)

    for row_idx, (label, value) in enumerate(estimate_header_rows(meta), start=2):
        cells = header_table.rows[row_idx].cells
        left = cells[0].merge(cells[1])
        right = cells[3].merge(cells[4])
        set_cell(left, label)
        set_cell(cells[2], "")
        set_cell(right, str(value or ""), bold=bool(value))

    clear_table_borders(header_table)
    document.add_paragraph()

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["№", "Вид работ", "Источник", "Показатель", "Расчет", "Коэффициенты", "Стоимость, тг"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell(cell, header, bold=True, center=True, size=9)

    for index, item in enumerate(items, start=1):
        row = table.add_row().cells
        set_cell(row[0], str(index), center=True, size=9)
        set_cell(row[1], str(item.get("title", "")), size=9)
        set_cell(row[2], str(item.get("source", "")), size=9)
        set_cell(row[3], str(item.get("indicatorText", "")), size=9)
        set_cell(row[4], str(item.get("formula", "")), size=9)
        coefficient_text = item.get("coefficientText") or f"{float(item.get('coefficient', 1) or 1):.4g}"
        set_cell(row[5], str(coefficient_text), size=9)
        set_cell(row[6], f"{float(item.get('total', 0) or 0):,.2f}".replace(",", " "), center=True, size=9)

    total_paragraph = document.add_paragraph()
    total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_run = total_paragraph.add_run(f"Итого: {total:,.2f} тг".replace(",", " "))
    total_run.bold = True
    total_run.font.name = "Times New Roman"
    total_run.font.size = Pt(12)

    note = meta.get("note", "")
    if note:
        document.add_paragraph(f"Примечание: {note}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_xlsx(payload: dict[str, Any]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    meta = payload.get("meta", {})
    items = payload.get("items", [])

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Смета"

    thin = Side(style="thin", color="D6D6D6")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_font = Font(name="Times New Roman", bold=True, size=12)
    text_font = Font(name="Times New Roman", size=11)
    bold_font = Font(name="Times New Roman", bold=True, size=11)

    sheet.merge_cells("A2:G2")
    sheet["A2"] = estimate_title(meta)
    sheet["A2"].font = header_font
    sheet["A2"].alignment = Alignment(horizontal="center", vertical="center")
    sheet.merge_cells("A3:G3")
    sheet["A3"] = estimate_subtitle(meta)
    sheet["A3"].font = header_font
    sheet["A3"].alignment = Alignment(horizontal="center", vertical="center")

    for row_no, (label, value) in enumerate(estimate_header_rows(meta), start=5):
        sheet.merge_cells(start_row=row_no, start_column=1, end_row=row_no, end_column=4)
        sheet.merge_cells(start_row=row_no, start_column=5, end_row=row_no, end_column=7)
        label_cell = sheet.cell(row=row_no, column=1, value=label)
        value_cell = sheet.cell(row=row_no, column=5, value=value or "")
        label_cell.font = text_font
        value_cell.font = bold_font if value else text_font
        label_cell.alignment = Alignment(wrap_text=True, vertical="center")
        value_cell.alignment = Alignment(wrap_text=True, vertical="center")
        sheet.row_dimensions[row_no].height = 34 if row_no in (6, 8) else 28

    for row in sheet.iter_rows(min_row=2, max_row=8, min_col=1, max_col=7):
        for cell in row:
            cell.border = border

    headers = ["№", "Вид работ", "Источник", "Показатель", "Расчет", "Коэффициенты", "Стоимость, тг"]
    start_row = 11
    fill = PatternFill("solid", fgColor="DCE8E1")
    for col, header in enumerate(headers, start=1):
        cell = sheet.cell(row=start_row, column=col, value=header)
        cell.font = Font(name="Times New Roman", bold=True, size=10)
        cell.fill = fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for idx, item in enumerate(items, start=1):
        row_no = start_row + idx
        values = [
            idx,
            item.get("title", ""),
            item.get("source", ""),
            item.get("indicatorText", ""),
            item.get("formula", ""),
            item.get("coefficientText") or float(item.get("coefficient", 1) or 1),
            float(item.get("total", 0) or 0),
        ]
        for col, value in enumerate(values, start=1):
            cell = sheet.cell(row=row_no, column=col, value=value)
            cell.border = border
            cell.font = Font(name="Times New Roman", size=10)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        sheet.cell(row=row_no, column=7).number_format = '# ##0.00'

    total_row = start_row + len(items) + 2
    sheet.cell(row=total_row, column=6, value="Итого").font = Font(name="Times New Roman", bold=True, size=11)
    sheet.cell(row=total_row, column=7, value=f"=SUM(G{start_row + 1}:G{start_row + len(items)})")
    sheet.cell(row=total_row, column=7).font = Font(name="Times New Roman", bold=True, size=11)
    sheet.cell(row=total_row, column=7).number_format = '# ##0.00'

    note = meta.get("note", "")
    if note:
        sheet.cell(row=total_row + 2, column=1, value="Примечание").font = Font(name="Times New Roman", bold=True, size=10)
        sheet.cell(row=total_row + 2, column=2, value=note)
        sheet.merge_cells(start_row=total_row + 2, start_column=2, end_row=total_row + 2, end_column=7)

    widths = [5, 22, 24, 17, 28, 24, 18]
    for col, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(col)].width = width
    sheet.freeze_panes = "A9"
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


class EstimateHandler(SimpleHTTPRequestHandler):
    server_version = "EstimateServer/0.1"

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def log_message(self, format: str, *args: Any) -> None:
        sys.stdout.write("[%s] %s\n" % (datetime.now().strftime("%H:%M:%S"), format % args))

    def send_json(self, data: Any, status: int = 200) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_binary(self, body: bytes, content_type: str, filename: str) -> None:
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{filename}")
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path
        if path == "/api/catalog":
            self.send_json(build_catalog())
            return
        if path == "/":
            self.path = "/static/index.html"
            return super().do_GET()
        if path.startswith("/static/"):
            static_path = (ROOT / path.lstrip("/")).resolve()
            if not str(static_path).startswith(str(STATIC_DIR.resolve())) or not static_path.exists():
                self.send_error(404)
                return
            content_type = mimetypes.guess_type(static_path.name)[0] or "application/octet-stream"
            data = static_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        self.send_error(404)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        try:
            payload = read_payload(self)
            meta = payload.setdefault("meta", {})
            if not meta.get("date"):
                meta["date"] = date.today().isoformat()
            base_name = clean_filename(meta.get("estimateNo") or "smeta")
            stamp = datetime.now().strftime("%Y%m%d_%H%M")

            if parsed.path == "/api/export/docx":
                body = create_docx(payload)
                self.send_binary(
                    body,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    f"{base_name}_{stamp}.docx",
                )
                return
            if parsed.path == "/api/export/xlsx":
                body = create_xlsx(payload)
                self.send_binary(
                    body,
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    f"{base_name}_{stamp}.xlsx",
                )
                return

            self.send_error(404)
        except Exception as exc:
            self.send_json({"error": str(exc)}, status=500)


def serve() -> None:
    parse_topo_rates()
    for port in range(8765, 8795):
        try:
            server = ThreadingHTTPServer(("127.0.0.1", port), EstimateHandler)
            break
        except OSError:
            continue
    else:
        raise RuntimeError("Не удалось найти свободный порт в диапазоне 8765-8794")

    print(f"Сервер запущен: http://127.0.0.1:{server.server_port}")
    print("Для остановки нажмите Ctrl+C")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nСервер остановлен")


if __name__ == "__main__":
    serve()
